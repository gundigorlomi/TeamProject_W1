"""Generate the Veracity presentation speech as a .docx ready to upload to
Google Docs (File → Open → Upload).

Three speakers, ~6 minutes total. Each block is one slide.
"""
from __future__ import annotations

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


SPEAKERS = ["Speaker A", "Speaker B", "Speaker C"]


# (slide #, slide title, speaker index, time, lines)
SCRIPT = [
    (1, "Cover — Veracity", 0, "0:25", [
        "Hi everyone. We're walking you through Veracity — a tool we built to "
        "answer one question fast: is this influencer's audience real?",
        "I'll start with the problem and the idea. My teammates will pick up the "
        "business side and the technical side after that.",
    ]),
    (2, "Project Overview", 0, "0:30", [
        "Here's the whole product in one screen. You paste a handle, and within "
        "a few seconds you get a single score from zero to a hundred.",
        "Underneath that score sit four signals — follower quality, engagement, "
        "audience fit, and growth — plus a 90-day timeline, a list of red flags, "
        "and a sample of suspicious followers. One report. One verdict. The "
        "evidence right next to it.",
    ]),
    (3, "Project Purpose", 0, "0:30", [
        "Why we built this. Today, vetting a creator is a manual job — "
        "screenshots, spreadsheets, gut feel. It eats an hour and the answer "
        "is still subjective.",
        "Our goal was simple. Take that hour and compress it into a glance. "
        "Same evidence, faster decision, and a number every stakeholder agrees on.",
    ]),
    (4, "The Problem", 0, "0:35", [
        "The problem we're solving is real money. Roughly half the engagement "
        "on big influencer accounts is artificial — bought followers, bot "
        "comments, engagement pods.",
        "Brands lose somewhere around 1.3 billion dollars a year to this. And "
        "by the time a campaign underperforms, the budget is already gone. There "
        "is no industry-standard score that catches it before contracts get signed.",
    ]),
    (5, "What It Solves", 0, "0:30", [
        "Veracity gives you that score. Four signals, one verdict, and the "
        "drill-downs are right there if anyone asks why.",
        "And it does this without anyone needing to be a data scientist. The "
        "same report works for an agency intern, a brand manager, and a "
        "legal team."
    ]),
    (6, "Core Idea", 1, "0:35", [
        "Thanks. I'll pick up the idea behind the engine.",
        "The core trick is this — every signal is a small, pure function. One "
        "function looks at the followers themselves. One looks at the shape of "
        "engagement over time. One checks whether the audience actually matches "
        "the niche. One looks at growth velocity for unnatural spikes.",
        "Each function is testable on its own. The composer at the top weighs "
        "them and produces the final score. So when we want to add a new check "
        "later — for example, comment timing patterns — it's one file, not a rewrite."
    ]),
    (7, "Target Audience", 1, "0:30", [
        "Who pays for this. The closer a team is to the spend, the more they "
        "care about the answer.",
        "Influencer agencies vetting their roster. Direct-to-consumer marketing "
        "teams writing the checks. PR teams protecting reputation. And talent "
        "platforms that need to score creators at sign-up. Mostly small to "
        "mid-size B2B — five to two hundred seats.",
    ]),
    (8, "Value Proposition", 1, "0:30", [
        "The pitch is straightforward. Veracity pays for itself the first "
        "time it stops a bad deal.",
        "An hour of guesswork becomes a ten-second answer. You save spend, you "
        "save time, and you have receipts for every creator you sign — which "
        "matters when something goes wrong later.",
    ]),
    (9, "Growth Strategy", 1, "0:30", [
        "Growth-wise, we're going land-and-expand. Anyone can run a free single-"
        "handle scan and share the result — that's the acquisition loop.",
        "From there we move them onto a team plan with bulk scanning, watchlists, "
        "and exports. The next layer is an API tier that drops scores straight "
        "into talent CRMs. Every scan we run also sharpens the model — so the "
        "longer customers stay, the harder we are to leave.",
    ]),
    (10, "Revenue Model", 1, "0:30", [
        "Revenue is recurring SaaS with a usage ceiling.",
        "Free plan for five scans a month. Pro at 49 dollars a month for two "
        "hundred scans. Team at 199 dollars for a thousand scans, with multi-"
        "seat and audit logs. And the API tier charges five cents per scan once "
        "you go past your plan. We're targeting around 85% gross margin, which "
        "the stack supports because hosting cost is small and most of the cost "
        "is bounded API calls.",
    ]),
    (11, "Costs & Sustainability", 2, "0:35", [
        "I'll take it from here on the technical and operational side.",
        "On costs — the stack is intentionally cheap. At MVP scale we're looking "
        "at about 20 to 80 dollars a month for hosting. Postgres adds pennies "
        "per gigabyte. The variable cost lives in the platform API calls, "
        "and we cap that two ways: per-plan ceilings on the customer side, and "
        "aggressive caching on ours — repeat scans for the same handle within "
        "a fresh window reuse the data instead of paying for new calls.",
    ]),
    (12, "Growth & Business Impact", 2, "0:30", [
        "On the business side — Year 1 we're targeting 200 paying teams and "
        "around a quarter million in ARR. Year 2 we want the API to be embedded "
        "in three or more talent CRMs, pushing past a million in ARR.",
        "Our defensibility comes from scan history. Once a team has a year of "
        "vetted creators in their workspace, ripping that out is painful. "
        "That's the moat.",
    ]),
    (13, "Technical Structure", 2, "0:40", [
        "Three layers, every one of them swappable on its own.",
        "On top, a React and TypeScript app built with Vite — the entire "
        "dashboard, the compare view, the live charts. In the middle, a "
        "FastAPI service with JWT auth, and scans that run in background tasks "
        "so the UI doesn't block.",
        "Underneath that sits the scoring engine — four pure functions, "
        "weights tunable in one place. And the data layer goes through "
        "a Provider Protocol. Right now we have a deterministic mock provider; "
        "swapping it for the real Instagram or TikTok API doesn't touch the "
        "engine, the API, or the UI. That's the part we're proud of.",
    ]),
    (14, "System Features", 2, "0:30", [
        "On features — the dashboard is the heart of it. Animated gauge, "
        "subscores that fill in, and a 90-day, 6-month, or 1-year engagement "
        "timeline you can toggle through.",
        "Compare mode lets you put two creators side by side. Red flags are "
        "ranked by severity. And we sample the suspicious followers themselves "
        "with a reason for each one — so it's not just 'these are bad,' it's "
        "'this account has no avatar, no posts, and follows 7,000 people.'",
    ]),
    (15, "Security & Ethics", 2, "0:35", [
        "Last slide — security and ethics, because a scoring product is "
        "only as good as its accountability.",
        "Passwords are bcrypt-hashed. Tokens are short-lived JWTs, no server-"
        "side session sprawl. We only ever look at public data — no DMs, no "
        "scraped private accounts. And every red flag is explainable. Nothing "
        "is a black box. If a creator gets a low score, we can show exactly which "
        "signals dragged it down.",
        "That's Veracity. Happy to take questions."
    ]),
]


def add_h1(doc, text, color=(15, 23, 42)):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(22)
    r.font.bold = True
    r.font.color.rgb = RGBColor(*color)
    return p


def add_h2(doc, text, color=(67, 56, 202)):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(13)
    r.font.bold = True
    r.font.color.rgb = RGBColor(*color)
    return p


def add_meta(doc, text, color=(100, 116, 139)):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(*color)
    r.italic = True
    return p


def add_body(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(30, 41, 59)
    p.paragraph_format.space_after = Pt(6)
    return p


def add_rule(doc):
    p = doc.add_paragraph()
    r = p.add_run("─" * 60)
    r.font.size = Pt(8)
    r.font.color.rgb = RGBColor(203, 213, 225)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def main():
    doc = Document()

    # page margins
    for section in doc.sections:
        section.left_margin = Inches(0.9)
        section.right_margin = Inches(0.9)
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)

    # cover header
    add_h1(doc, "Veracity — Presentation Speech")
    add_meta(doc, "Influencer Authenticity Checker  ·  ~6 minutes  ·  3 speakers")

    # speaker map
    add_h2(doc, "Speakers")
    speaker_blocks = {0: [], 1: [], 2: []}
    for s in SCRIPT:
        speaker_blocks[s[2]].append(f"Slide {s[0]} ({s[3]}) — {s[1]}")
    for i, who in enumerate(SPEAKERS):
        p = doc.add_paragraph()
        r = p.add_run(f"{who}: ")
        r.font.bold = True
        r.font.size = Pt(11)
        r.font.color.rgb = RGBColor(30, 41, 59)
        r2 = p.add_run("; ".join(speaker_blocks[i]))
        r2.font.size = Pt(10)
        r2.font.color.rgb = RGBColor(71, 85, 105)

    add_rule(doc)

    # tips
    add_h2(doc, "Delivery notes")
    tips = [
        "Pace: about 145 words a minute. The whole script clocks in around six minutes.",
        "Hand-offs: end your last slide with 'I'll let [name] take it from here' — the script does this naturally on slides 6 and 11.",
        "If a slide animates in pieces (the bullet cards fade in one by one), pause briefly between sentences so the visuals catch up.",
        "Don't read line by line — these are talking points. Glance at the slide, then the audience.",
    ]
    for t in tips:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(t)
        r.font.size = Pt(10)
        r.font.color.rgb = RGBColor(51, 65, 85)

    add_rule(doc)

    # script body
    add_h1(doc, "Script", color=(15, 23, 42))
    for slide_no, title, speaker_idx, runtime, lines in SCRIPT:
        # slide label
        p = doc.add_paragraph()
        r = p.add_run(f"SLIDE {slide_no:02d}  ·  {title}")
        r.font.size = Pt(12)
        r.font.bold = True
        r.font.color.rgb = RGBColor(15, 23, 42)
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(2)

        # who + runtime
        p = doc.add_paragraph()
        r = p.add_run(f"{SPEAKERS[speaker_idx]}  ·  {runtime}")
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(99, 102, 241)
        r.italic = True
        p.paragraph_format.space_after = Pt(4)

        for line in lines:
            add_body(doc, line)

    add_rule(doc)
    add_meta(doc, "End of script.")

    out = "Veracity_Speech.docx"
    doc.save(out)
    print(f"Wrote {out}.")


if __name__ == "__main__":
    main()
